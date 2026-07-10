"""CRUD de modelos de petição reutilizáveis (tenant-scoped)."""
from fastapi import APIRouter, Depends, Query, UploadFile, File, Form, HTTPException
from fastapi.responses import Response
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, desc
from pydantic import BaseModel
import uuid

from app.db.base import get_db
from app.dependencies import get_current_user
from app.models.user import User
from app.models.document import PetitionTemplate
from app.core.exceptions import NotFoundError

router = APIRouter(prefix="/petition-templates", tags=["petition-templates"])


class TemplateResponse(BaseModel):
    id: str
    nome: str
    tipo_peticao: str | None
    descricao: str | None
    conteudo: str
    ativo: bool
    created_at: str


class TemplateCreate(BaseModel):
    nome: str
    tipo_peticao: str | None = None
    descricao: str | None = None
    conteudo: str
    ativo: bool = True


class TemplateUpdate(BaseModel):
    nome: str | None = None
    tipo_peticao: str | None = None
    descricao: str | None = None
    conteudo: str | None = None
    ativo: bool | None = None


def _to_response(t: PetitionTemplate) -> TemplateResponse:
    return TemplateResponse(
        id=str(t.id),
        nome=t.nome,
        tipo_peticao=t.tipo_peticao,
        descricao=t.descricao,
        conteudo=t.conteudo,
        ativo=t.ativo,
        created_at=t.created_at.isoformat(),
    )


@router.get("", response_model=list[TemplateResponse])
async def list_templates(
    tipo_peticao: str | None = None,
    ativo: bool | None = None,
    limit: int = Query(default=100, le=200),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    query = (
        select(PetitionTemplate)
        .where(PetitionTemplate.tenant_id == current_user.tenant_id)
        .order_by(desc(PetitionTemplate.created_at))
        .limit(limit)
    )
    if tipo_peticao:
        query = query.where(PetitionTemplate.tipo_peticao == tipo_peticao)
    if ativo is not None:
        query = query.where(PetitionTemplate.ativo == ativo)
    result = await db.execute(query)
    return [_to_response(t) for t in result.scalars().all()]


@router.post("", status_code=201, response_model=TemplateResponse)
async def create_template(
    body: TemplateCreate,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    tpl = PetitionTemplate(
        nome=body.nome,
        tipo_peticao=body.tipo_peticao,
        descricao=body.descricao,
        conteudo=body.conteudo,
        ativo=body.ativo,
        created_by=current_user.id,
        tenant_id=current_user.tenant_id,
    )
    db.add(tpl)
    await db.flush()
    return _to_response(tpl)


async def _get_owned(db: AsyncSession, template_id: str, current_user: User) -> PetitionTemplate:
    result = await db.execute(
        select(PetitionTemplate).where(
            PetitionTemplate.id == uuid.UUID(template_id),
            PetitionTemplate.tenant_id == current_user.tenant_id,
        )
    )
    tpl = result.scalar_one_or_none()
    if not tpl:
        raise NotFoundError("PetitionTemplate", template_id)
    return tpl


@router.put("/{template_id}", response_model=TemplateResponse)
async def update_template(
    template_id: str,
    body: TemplateUpdate,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    tpl = await _get_owned(db, template_id, current_user)
    for field, value in body.model_dump(exclude_none=True).items():
        setattr(tpl, field, value)
    await db.flush()
    return _to_response(tpl)


@router.delete("/{template_id}", status_code=204)
async def delete_template(
    template_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    tpl = await _get_owned(db, template_id, current_user)
    await db.delete(tpl)
    await db.flush()


def _extract_docx_text(raw: bytes) -> str:
    """Extrai o texto de um .docx (parágrafos + tabelas), preservando quebras."""
    import io
    from docx import Document as DocxDocument

    docx = DocxDocument(io.BytesIO(raw))
    lines: list[str] = []
    for p in docx.paragraphs:
        lines.append(p.text)
    for table in docx.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines).strip()


@router.post("/upload", status_code=201, response_model=TemplateResponse)
async def upload_template(
    file: UploadFile = File(...),
    nome: str | None = Form(None),
    tipo_peticao: str | None = Form(None),
    descricao: str | None = Form(None),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Importa um modelo pronto do Word (.docx) ou texto (.txt).

    O conteúdo extraído fica editável no app e pode ser baixado de volta em
    .docx para edição no Word e reimportação."""
    raw = await file.read()
    if not raw:
        raise HTTPException(status_code=400, detail="Arquivo vazio.")
    if len(raw) > 5 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Arquivo muito grande. Máximo 5MB.")

    filename = file.filename or "modelo"
    lower = filename.lower()
    if lower.endswith(".docx"):
        try:
            conteudo = _extract_docx_text(raw)
        except Exception:
            raise HTTPException(status_code=400, detail="Não foi possível ler o .docx. Verifique o arquivo.")
    elif lower.endswith(".txt"):
        conteudo = raw.decode("utf-8", errors="ignore").strip()
    else:
        raise HTTPException(status_code=400, detail="Formato não suportado. Envie .docx ou .txt.")

    if not conteudo:
        raise HTTPException(status_code=400, detail="O arquivo não contém texto extraível.")

    tpl = PetitionTemplate(
        nome=(nome or filename.rsplit(".", 1)[0])[:200],
        tipo_peticao=tipo_peticao,
        descricao=descricao,
        conteudo=conteudo,
        ativo=True,
        created_by=current_user.id,
        tenant_id=current_user.tenant_id,
    )
    db.add(tpl)
    await db.flush()
    return _to_response(tpl)


@router.get("/{template_id}/docx")
async def download_template_docx(
    template_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Baixa o modelo como .docx para edição no Word (depois reimporte)."""
    import io
    from docx import Document as DocxDocument

    tpl = await _get_owned(db, template_id, current_user)
    docx = DocxDocument()
    for line in tpl.conteudo.split("\n"):
        docx.add_paragraph(line)
    buf = io.BytesIO()
    docx.save(buf)
    safe_name = "".join(c for c in tpl.nome if c.isalnum() or c in " _-")[:50] or "modelo"
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}.docx"'},
    )
