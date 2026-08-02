"""Extração de texto de arquivos .docx (parágrafos + tabelas)."""
import io


def extract_docx_text(raw: bytes) -> str:
    """Extrai o texto de um .docx (parágrafos + tabelas), preservando quebras."""
    from docx import Document as DocxDocument

    docx = DocxDocument(io.BytesIO(raw))
    lines: list[str] = []
    for p in docx.paragraphs:
        lines.append(p.text)
    for table in docx.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines).strip()
